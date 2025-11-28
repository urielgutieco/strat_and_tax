import os
import io
import zipfile
import random
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Inches

# Nuevas importaciones para Google Drive
import pickle
import json
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

app = Flask(__name__)
CORS(app)

# Carpetas raíz
TEMPLATE_FOLDER = 'template_word'
GENERATED_DOCS = 'template_gendocs'
GENERATED_ZIPS = 'template_genzips'

# Si modificas el alcance de acceso (SCOPE), borra el archivo token.pickle
SCOPES = ['https://www.googleapis.com/auth/drive.file']
TOKEN_FILE = 'token.pickle'
CREDENTIALS_FILE = 'credentials.json'

os.makedirs(TEMPLATE_FOLDER, exist_ok=True)
os.makedirs(GENERATED_DOCS, exist_ok=True)
os.makedirs(GENERATED_ZIPS, exist_ok=True)

# Mapeo: valor del select -> subcarpeta en template_word
SERVICIO_TO_DIR = {
    "Servicios de construccion de unidades unifamiliares": "construccion_unifamiliar",
    "Servicios de reparacion o ampliacion o remodelacion de viviendas unifamiliares": "reparacion_remodelacion_unifamiliar",
    "Servicio de remodelacion general de viviendas unifamiliares": "remodelacion_general",
    "Servicios de reparacion de casas moviles en el sitio": "reparacion_casas_moviles",
    "Servicios de construccion y reparacion de patios y terrazas": "patios_terrazas",
    "Servico de reparacion por daños ocasionados por fuego de viviendas unifamiliares": "reparacion_por_fuego",
    "Servicio de construccion de casas unifamiliares nuevas": "construccion_unifamiliar_nueva",
    "Servicio de instalacion de casas unifamiliares prefabricadas": "instalacion_prefabricadas",
    "Servicio de construccion de casas en la ciudad o casas jardin unifamiliares nuevas": "construccion_casas_ciudad_jardin",
    "Dasarrollo urbano": "desarrollo_urbano",  # coincide con el valor tal como está escrito
    "Servicio de planificacion de la ordenacion urbana": "planificacion_ordenacion_urbana",
    "Servicio de administracion de tierras urbanas": "administracion_tierras_urbanas",
    "Servicio de programacion de inversiones urbanas": "programacion_inversiones_urbanas",
    "Servicio de reestructuracion de barrios marginales": "reestructuracion_barrios_marginales",
    "Servicios de alumbrado urbano": "alumbrado_urbano",
    "Servicios de control o regulacion del desarrollo urbano": "control_desarrollo_urbano",
    "Servicios de estandares o regulacion de edificios urbanos": "estandares_regulacion_edificios",
    "Servicios comunitarios urbanos": "comunitarios_urbanos",
    "Servicios de administracion o gestion de proyectos o programas urbanos": "gestion_proyectos_programas_urbanos",
    "Ingenieria civil": "ingenieria_civil",
    "Ingenieria de carreteras": "ingenieria_carreteras",
    "Ingenieria deinfraestructura de instalaciones o fabricas": "infraestructura_instalaciones_fabricas",
    "Servicios de mantenimiento e instalacion de equipo pesado": "mantenimiento_instalacion_equipo_pesado",
    "Servicio de mantenimiento y reparacion de equipo pesado": "mantenimiento_reparacion_equipo_pesado",
}

# Plantillas (nombres comunes existentes en cada subcarpeta)
TEMPLATE_FILES = [
    'plantilla_solicitud.docx',
    '2.docx',
    '3.docx',
    '4.docx',
    '1.docx',
]

def replace_text_in_document(document, replacements):
    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))

def generate_single_document(template_filename, template_root, replacements, user_image_path=None, data=None):
    template_path = os.path.join(template_root, template_filename)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Plantilla '{template_filename}' no encontrada en '{template_root}'.")

    document = Document(template_path)
    replace_text_in_document(document, replacements)

    # Imagen del usuario (opcional)
    if user_image_path and os.path.exists(user_image_path):
        try:
            # Es necesario asegurar que los párrafos se inserten en el lugar deseado
            # Aquí se inserta la imagen al final del documento por defecto
            document.add_paragraph()
            document.add_paragraph(data.get('nombre_completo_de_la_persona_que_firma_la_solicitud', 'N/A') if data else 'N/A')
            document.add_picture(user_image_path, width=Inches(2.5))
        except Exception:
            document.add_paragraph("⚠ No se pudo insertar la imagen del usuario.")
    else:
        document.add_paragraph("⚠ Imagen de firma no encontrada en el servidor.")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# Función ADICIONAL para manejar la autenticación y la carga a Drive
def authenticate_and_upload_to_drive(file_name, zip_buffer):
    """
    Realiza la autenticación OAuth2 y sube el archivo ZIP a Google Drive.
    """
    creds = None
    # 1. Cargar credenciales guardadas
    if os.path.exists(TOKEN_FILE):
        with open(TOKEN_FILE, 'rb') as token:
            creds = pickle.load(token)
            
    # 2. Manejar credenciales (refrescar o iniciar flujo)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            # Asegúrate de tener tu archivo credentials.json de Google Cloud Console
            if not os.path.exists(CREDENTIALS_FILE):
                print(f"Error: No se encontró el archivo de credenciales '{CREDENTIALS_FILE}'.")
                return {"success": False, "message": "Falta el archivo de credenciales de Google Drive."}
            
            # Nota: Este flujo de InstalledAppFlow.run_local_server()
            # funciona solo si el servidor Flask se ejecuta localmente y
            # tiene acceso al navegador del usuario para el inicio de sesión.
            flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_FILE, SCOPES)
            creds = flow.run_local_server(port=0)
            
        # Guardar las credenciales para la próxima ejecución
        with open(TOKEN_FILE, 'wb') as token:
            pickle.dump(creds, token)

    try:
        service = build('drive', 'v3', credentials=creds)
        
        # Metadata del archivo
        file_metadata = {'name': file_name}
        
        # **IMPORTANTE:** Rebobinamos el buffer al inicio antes de la carga.
        zip_buffer.seek(0)
        
        # El MediaIoBaseUpload toma el objeto io.BytesIO como medio
        media = MediaIoBaseUpload(zip_buffer, mimetype='application/zip', resumable=True)
        
        # Llama a la API para cargar el archivo
        file = service.files().create(body=file_metadata,
                                      media_body=media,
                                      fields='id').execute()
        
        return {"success": True, "message": f"Archivo cargado a Google Drive con éxito. ID: {file.get('id')}"}

    except Exception as e:
        # Aquí puedes manejar errores específicos de la API (p.ej. cuota)
        return {"success": False, "message": f"Error al cargar a Google Drive: {str(e)}"}


@app.route('/generate-word', methods=['POST'])
def generate_word():
    try:
        data = request.form.to_dict()

        uploaded_image = request.files.get("imagen_usuario")
        user_image_path = None
        if uploaded_image:
            # Guardamos la imagen temporalmente para que docx la pueda leer
            user_image_path = os.path.join(TEMPLATE_FOLDER, "imagen_custom.png")
            uploaded_image.save(user_image_path)

        if not data:
            return jsonify({"error": "No data received"}), 400

        # Limpieza de generated_docs (opcional, se mantiene del código original)
        for filename in os.listdir(GENERATED_DOCS):
            file_path = os.path.join(GENERATED_DOCS, filename)
            try:
                if os.path.isfile(file_path):
                    os.remove(file_path)
            except Exception as e:
                print(f"Error al eliminar el archivo {file_path}: {e}")

        # Servicio seleccionado y carpeta asociada
        servicio = data.get('servicio')
        if not servicio:
            return jsonify({"error": "Debes seleccionar un servicio."}), 400

        carpeta_servicio = SERVICIO_TO_DIR.get(servicio)
        if not carpeta_servicio:
            return jsonify({"error": f"No existe carpeta mapeada para el servicio: {servicio}"}), 404

        # Root dinámico de plantillas según servicio
        template_root = os.path.join(TEMPLATE_FOLDER, carpeta_servicio)
        if not os.path.isdir(template_root):
            return jsonify({"error": f"La carpeta de plantillas no existe: {template_root}"}), 404

        # Número de contrato único
        numero_de_contrato_unico = ''.join([str(random.randint(0, 9)) for _ in range(18)])

        # Forzar la descripción del servicio desde el select
        descripcion_servicio = servicio

        replacements = {
            '${descripcion_del_servicio}': descripcion_servicio,
            '${razon_social}': data.get('razon_social', 'N/A'),
            '${r_f_c}': data.get('r_f_c', 'N/A'),
            '${domicilio_del_cliente}': data.get('domicilio_del_cliente', 'N/A'),
            '${telefono_del__cliente}': data.get('telefono_del__cliente', 'N/A'),
            '${correo_del_cliente}': data.get('correo_del_cliente', 'N/A'),
            '${fecha_de_inicio_del_servicio}': data.get('fecha_de_inicio_del_servicio', 'N/A'),
            '${fecha_de_conclusion_del_servicio}': data.get('fecha_de_conclusion_del_servicio', 'N/A'),
            '${monto_de_la_operacion_Sin_IVA}': data.get('monto_de_la_operacion_Sin_IVA', 'N/A'),
            '${forma_de_pago}': data.get('forma_de_pago', 'N/A'),
            '${cantidad}': data.get('cantidad', 'N/A'),
            '${unidad}': data.get('unidad', 'N/A'),
            '${numero_de_contrato}': numero_de_contrato_unico,
            '${fecha_de_operación}': data.get('fecha_de_operación', 'N/A'),
            '${nombre_completo_de_la_persona_que_firma_la_solicitud}': data.get('nombre_completo_de_la_persona_que_firma_la_solicitud', 'N/A'),
            '${cargo_de_la_persona_que_firma_la_solicitud}': data.get('cargo_de_la_persona_que_firma_la_solicitud', 'N/A'),
            '${factura_relacionada_con_la_operación}': data.get('factura_relacionada_con_la_operación', 'N/A'),
            '${informe_si_cuenta_con_fotografias_videos_o_informacion_adicion}': data.get('informe_si_cuenta_con_fotografias_videos_o_informacion_adicion', 'N/A'),
            '${comentarios}': data.get('comentarios', 'N/A')
        }

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for template in TEMPLATE_FILES:
                try:
                    doc_buffer = generate_single_document(template, template_root, replacements, user_image_path, data)
                    base = os.path.splitext(template)[0]

                    # Nombre de salida legible
                    rfc = data.get('r_f_c', 'N/A')
                    output_filename = f"{base}_{descripcion_servicio}_{base}_{numero_de_contrato_unico}_{rfc}.docx"

                    output_path = os.path.join(GENERATED_DOCS, output_filename)
                    if os.path.exists(output_path):
                        os.remove(output_path)
                    
                    with open(output_path, "wb") as f:
                        f.write(doc_buffer.getvalue())

                    zip_file.writestr(output_filename, doc_buffer.getvalue())

                except Exception as e:
                    print(f"Error generando documento {template}: {e}")
                    continue

        zip_buffer.seek(0)
        final_zip_name = f"{descripcion_servicio}_{numero_de_contrato_unico}_{data.get('r_f_c', 'N/A')}.zip"

        # Guardado local del ZIP
        zip_server_path = os.path.join(GENERATED_ZIPS, final_zip_name)
        with open(zip_server_path, "wb") as zip_file_for_storage:
            zip_file_for_storage.write(zip_buffer.getvalue())
        
        # Lógica ADICIONAL: Llamar a la función para subir el archivo a Google Drive
        # Nota: Volvemos a posicionar el buffer al inicio antes de la carga a Drive
        zip_buffer.seek(0) 
        upload_result = authenticate_and_upload_to_drive(final_zip_name, zip_buffer)
        
        # El mensaje de Google Drive se imprimirá en la consola del servidor (terminal)
        print(f"Resultado de la carga a Google Drive: {upload_result['message']}")
        
        # Volvemos a posicionar el buffer al inicio antes de enviarlo al cliente (descarga)
        zip_buffer.seek(0) 
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=final_zip_name
        )

    except Exception as e:
        return jsonify({"error": f"Error interno: {str(e)}"}), 500

if __name__ == '__main__':
    try:
        # Nota: Para que la autenticación de Drive funcione,
        # necesitarás tener el archivo 'credentials.json' en el mismo directorio
        # y la primera ejecución abrirá una ventana del navegador para iniciar sesión.
        app.run(debug=True, port=5001)
    except Exception as e:
        print(f"Error al iniciar la aplicación: {e}")