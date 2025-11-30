import os
import io
import zipfile
import random
import json
import base64
import logging
import tempfile
import pathlib
import re
import bcrypt
import jwt
import time
from flask import Flask, request, send_file, jsonify, abort
from flask_cors import CORS
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from datetime import datetime, timedelta, timezone
from functools import wraps
from collections import defaultdict

# -------------------------
# Configuración Flask / Producción
# -------------------------
app = Flask(__name__)

MAX_CONTENT_LENGTH_BYTES = int(os.getenv("MAX_CONTENT_LENGTH_BYTES", 20 * 1024 * 1024))
ALLOWED_IMAGE_EXT = {".png", ".jpg", ".jpeg"}
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH_BYTES
app.config["JSON_SORT_KEYS"] = False

allowed_origins = os.getenv("ALLOWED_ORIGINS", "")
if allowed_origins:
    origins = [o.strip() for o in allowed_origins.split(",") if o.strip()]
    CORS(app, origins=origins)
else:
    CORS(app, resources={r"/generate-word": {"origins": []}})

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("generate-word-app")

# -------------------------
# JWT Config
# -------------------------
JWT_SECRET_KEY = os.getenv("$2b$12$tUa5Z8rF.E.q2H/i.q5U7.G.R6A1W9V4P3I0Y2X5Q8T7S6R5C41")
if not JWT_SECRET_KEY:
    raise EnvironmentError("JWT_SECRET_KEY no configurada. Configúrala en Vercel.")
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_HOURS = 1  # Token expira en 1 hora

# -------------------------
# Usuarios controlados (solo usuarios normales)
# -------------------------
USER_DB = {
    "usuario1": {
        "password_hash": b"$2b$12$hGq9p0aQ4w7xS2zV.B.c.8A.D.E.F3G4H5I6J7K8L9M0N1O2P3Q4",
        "role": "user"
    },
    "usuario2": {
        "password_hash": b"$2b$12$Xyz123Abc456Def789Ghi012Jkl345Mno678Pqr901Stu234Vwx",
        "role": "user"
    }
}

# -------------------------
# Rate limiting básico por IP (login)
# -------------------------
MAX_ATTEMPTS = 5
BLOCK_TIME_SECONDS = 300  # 5 minutos
login_attempts = defaultdict(lambda: {"count": 0, "last_attempt": 0, "blocked_until": 0})

def check_rate_limit(ip):
    entry = login_attempts[ip]
    now = time.time()
    if entry["blocked_until"] > now:
        return False, int(entry["blocked_until"] - now)
    return True, 0

def record_failed_attempt(ip):
    entry = login_attempts[ip]
    now = time.time()
    entry["count"] += 1
    entry["last_attempt"] = now
    if entry["count"] >= MAX_ATTEMPTS:
        entry["blocked_until"] = now + BLOCK_TIME_SECONDS

def reset_attempts(ip):
    login_attempts[ip] = {"count": 0, "last_attempt": 0, "blocked_until": 0}

# -------------------------
# Carpetas y Mapeos
# -------------------------
BASE_DIR = pathlib.Path(os.getenv("BASE_DIR", "."))
TEMPLATE_FOLDER_NAME = os.getenv("TEMPLATE_FOLDER_NAME", "template_word")
TEMPLATE_FOLDER = pathlib.Path(__file__).parent / TEMPLATE_FOLDER_NAME
GENERATED_DOCS = BASE_DIR / 'template_gendocs'
GENERATED_ZIPS = BASE_DIR / 'template_genzips'

TEMPLATE_FOLDER.mkdir(parents=True, exist_ok=True)
GENERATED_DOCS.mkdir(parents=True, exist_ok=True)
GENERATED_ZIPS.mkdir(parents=True, exist_ok=True)

SERVICIO_TO_DIR = {
    "Servicios de construccion de unidades unifamiliares": "construccion_unifamiliar",
    "Servicios de reparacion o ampliacion o remodelacion de viviendas unifamiliares": "reparacion_remodelacion_unifamiliar",
    "Servicio de remodelacion general de viviendas unifamiliares": "remodelacion_general",
    "Servicios de reparacion de casas moviles en el sitio": "reparacion_casas_moviles",
    "Servicios de construccion y reparacion de patios y terrazas": "patios_terrazas",
    "Servico de reparacion por daños ocasionados por fuego de viviendas unifamiliares": "reparacion_por_fuego",
    "Servicio de construccion de casas unifamiliares nuevas": "construccion_unifamiliar_nueva",
    "Servicio de instalacion de casas unifamiliares prefabricadas": "instalacion_prefabricadas",
    "Servicio de construccion de casas en la ciudad o casas jardin unifamiliares nuevas": "construccion_casas_ciudad_jardin",
    "Dasarrollo urbano": "desarrollo_urbano",
    "Servicio de planificacion de la ordenacion urbana": "planificacion_ordenacion_urbana",
    "Servicio de administracion de tierras urbanas": "administracion_tierras_urbanas",
    "Servicio de programacion de inversiones urbanas": "programacion_inversiones_urbanas",
    "Servicio de reestructuracion de barrios marginales": "reestructuracion_barrios_marginales",
    "Servicios de alumbrado urbano": "alumbrado_urbano",
    "Servicios de control o regulacion del desarrollo urbano": "control_desarrollo_urbano",
    "Servicios de estandares o regulacion de edificios urbanos": "estandares_regulacion_edificios",
    "Servicios comunitarios urbanos": "comunitarios_urbanos",
    "Servicios de administracion o gestion de proyectos o programas urbanos": "gestion_proyectos_programas_urbanos",
    "Ingenieria civil": "ingenieria_civil",
    "Ingenieria de carreteras": "ingenieria_carreteras",
    "Ingenieria deinfraestructura de instalaciones o fabricas": "infraestructura_instalaciones_fabricas",
    "Servicios de mantenimiento e instalacion de equipo pesado": "mantenimiento_instalacion_equipo_pesado",
    "Servicio de mantenimiento y reparacion de equipo pesado": "mantenimiento_reparacion_equipo_pesado",
}

TEMPLATE_FILES = [
    'plantilla_solicitud.docx',
    '2.docx',
    '3.docx',
    '4.docx',
    '1.docx',
]

# -------------------------
# Utilidades
# -------------------------
def sanitize_filename(name: str) -> str:
    name = secure_filename(name)
    name = re.sub(r'[_]{2,}', '_', name)
    return name[:200]

def replace_text_in_document(document, replacements):
    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))

def generate_single_document(template_filename, template_root, replacements, user_image_path=None, data=None):
    template_path = os.path.join(template_root, template_filename)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Plantilla '{template_filename}' no encontrada en '{template_root}'.")
    document = Document(template_path)
    replace_text_in_document(document, replacements)
    if user_image_path and os.path.exists(user_image_path):
        try:
            document.add_paragraph()
            document.add_paragraph(data.get('nombre_completo_de_la_persona_que_firma_la_solicitud', 'N/A') if data else 'N/A')
            document.add_picture(user_image_path, width=Inches(2.5))
        except Exception as ex:
            logger.warning("No se pudo insertar la imagen del usuario: %s", ex)
            document.add_paragraph("⚠ No se pudo insertar la imagen del usuario.")
    else:
        document.add_paragraph("⚠ Imagen de firma no encontrada en el servidor.")
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# -------------------------
# Google Drive Service Account
# -------------------------
def _load_service_account_info():
    raw_json = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
    if raw_json:
        try:
            return json.loads(raw_json)
        except Exception as e:
            logger.error("GOOGLE_SERVICE_ACCOUNT_JSON no es JSON válido: %s", e)
            raise
    b64 = os.getenv("GOOGLE_SERVICE_ACCOUNT_BASE64")
    if b64:
        try:
            decoded = base64.b64decode(b64)
            return json.loads(decoded)
        except Exception as e:
            logger.error("GOOGLE_SERVICE_ACCOUNT_BASE64 inválido: %s", e)
            raise
    raise EnvironmentError("Falta GOOGLE_SERVICE_ACCOUNT_JSON o GOOGLE_SERVICE_ACCOUNT_BASE64")

def authenticate_and_upload_to_drive(file_name, zip_buffer):
    if os.getenv("DISABLE_DRIVE_UPLOAD", "0") == "1":
        logger.info("Subida a Drive deshabilitada por variable DISABLE_DRIVE_UPLOAD.")
        return {"success": True, "message": "Subida a Drive deshabilitada (env var)"}
    try:
        service_account_info = _load_service_account_info()
        creds = Credentials.from_service_account_info(service_account_info, scopes=["https://www.googleapis.com/auth/drive.file"])
        service = build('drive', 'v3', credentials=creds, cache_discovery=False)
        file_metadata = {'name': sanitize_filename(file_name)}
        zip_buffer.seek(0)
        media = MediaIoBaseUpload(zip_buffer, mimetype='application/zip', resumable=True)
        uploaded_file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
        return {"success": True, "message": f"Archivo subido correctamente. ID: {uploaded_file.get('id')}"}
    except Exception as e:
        logger.exception("Error al subir a Drive: %s", e)
        return {"success": False, "message": f"Error al subir a Drive: {str(e)}"}

# -------------------------
# Decorador JWT usuarios normales
# -------------------------
def token_required(required_role='user'):
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            token = None
            auth_header = request.headers.get('Authorization')
            if auth_header and auth_header.startswith('Bearer '):
                token = auth_header.split(' ')[1]
            if not token:
                return jsonify({'error': 'Token de acceso faltante. Inicia sesión.'}), 403
            try:
                data = jwt.decode(token, JWT_SECRET_KEY, algorithms=[JWT_ALGORITHM], audience='strat_and_tax_api', issuer='strat_and_tax_server')
                current_user = data['sub']
                role = data.get('role')
                if current_user not in USER_DB:
                    return jsonify({'error': 'Usuario no autorizado'}), 403
                if role != required_role:
                    return jsonify({'error': 'No tienes permisos para este recurso'}), 403
            except jwt.ExpiredSignatureError:
                return jsonify({'error': 'Token expirado. Vuelve a iniciar sesión.'}), 401
            except jwt.InvalidTokenError:
                return jsonify({'error': 'Token inválido o manipulado.'}), 403
            return f(current_user, *args, **kwargs)
        return decorated
    return decorator

# -------------------------
# Endpoint Login
# -------------------------
@app.route('/login', methods=['POST'])
def login():
    ip = request.remote_addr
    allowed, wait_time = check_rate_limit(ip)
    if not allowed:
        return jsonify({"error": f"Demasiados intentos. Espera {wait_time} segundos"}), 429
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')
    if not username or not password:
        return jsonify({"error": "Falta usuario o contraseña"}), 400
    user_record = USER_DB.get(username)
    if user_record and bcrypt.checkpw(password.encode('utf-8'), user_record['password_hash']):
        reset_attempts(ip)
        issued_at = datetime.now(timezone.utc)
        expiration_time = issued_at + timedelta(hours=JWT_EXPIRATION_HOURS)
        payload = {
            'sub': username,
            'role': user_record['role'],
            'iat': int(issued_at.timestamp()),
            'nbf': int(issued_at.timestamp()),
            'exp': int(expiration_time.timestamp()),
            'aud': 'strat_and_tax_api',
            'iss': 'strat_and_tax_server'
        }
        token = jwt.encode(payload, JWT_SECRET_KEY, algorithm=JWT_ALGORITHM)
        return jsonify({"message": "Login exitoso", "token": token}), 200
    else:
        record_failed_attempt(ip)
        return jsonify({"error": "Usuario o contraseña inválidos"}), 401

# -------------------------
# Endpoint principal protegido
# -------------------------
@app.route('/generate-word', methods=['POST'])
@token_required(required_role='user')
def generate_word(current_user):
    logger.info(f"Generación de documentos iniciada por usuario: {current_user}")
    try:
        data = request.form.to_dict()
        uploaded_image = request.files.get("imagen_usuario")
        user_image_path = None

        if not data:
            return jsonify({"error": "No data received"}), 400

        servicio = data.get('servicio')
        if not servicio:
            return jsonify({"error": "Debes seleccionar un servicio."}), 400
        carpeta_servicio = SERVICIO_TO_DIR.get(servicio)
        if not carpeta_servicio:
            return jsonify({"error": f"No existe carpeta mapeada para el servicio: {servicio}"}), 404
        template_root = os.path.join(TEMPLATE_FOLDER, carpeta_servicio)
        if not os.path.isdir(template_root):
            return jsonify({"error": f"La carpeta de plantillas no existe: {template_root}"}), 404

        if uploaded_image and uploaded_image.filename:
            filename = sanitize_filename(uploaded_image.filename)
            ext = pathlib.Path(filename).suffix.lower()
            if ext not in ALLOWED_IMAGE_EXT:
                return jsonify({"error": "Tipo de archivo no permitido para imagen."}), 400
            tmp_dir = tempfile.TemporaryDirectory()
            tmp_path = pathlib.Path(tmp_dir.name) / filename
            uploaded_image.save(str(tmp_path))
            user_image_path = str(tmp_path)

        # limpiar docs previos
        try:
            for filename in os.listdir(GENERATED_DOCS):
                file_path = os.path.join(GENERATED_DOCS, filename)
                if os.path.isfile(file_path):
                    try:
                        os.remove(file_path)
                    except Exception as e:
                        logger.warning("Error al eliminar %s: %s", file_path, e)
        except FileNotFoundError:
            pass

        numero_de_contrato_unico = ''.join([str(random.randint(0, 9)) for _ in range(18)])
        descripcion_servicio = servicio

        replacements = {
            '${descripcion_del_servicio}': descripcion_servicio,
            '${razon_social}': data.get('razon_social', 'N/A'),
            '${r_f_c}': data.get('r_f_c', 'N/A'),
            '${domicilio_del_cliente}': data.get('domicilio_del_cliente', 'N/A'),
            '${telefono_del__cliente}': data.get('telefono_del__cliente', 'N/A'),
            '${correo_del_cliente}': data.get('correo_del_cliente', 'N/A'),
            '${fecha_de_inicio_del_servicio}': data.get('fecha_de_inicio_del_servicio', 'N/A'),
            '${fecha_de_conclusion_del_servicio}': data.get('fecha_de_conclusion_del_servicio', 'N/A'),
            '${monto_de_la_operacion_Sin_IVA}': data.get('monto_de_la_operacion_Sin_IVA', 'N/A'),
            '${forma_de_pago}': data.get('forma_de_pago', 'N/A'),
            '${cantidad}': data.get('cantidad', 'N/A'),
            '${unidad}': data.get('unidad', 'N/A'),
            '${numero_de_contrato}': numero_de_contrato_unico,
            '${fecha_de_operación}': data.get('fecha_de_operacion', 'N/A'),
            '${nombre_completo_de_la_persona_que_firma_la_solicitud}': data.get('nombre_completo_de_la_persona_que_firma_la_solicitud', 'N/A'),
            '${cargo_de_la_persona_que_firma_la_solicitud}': data.get('cargo_de_la_persona_que_firma_la_solicitud', 'N/A'),
            '${factura_relacionada_con_la_operación}': data.get('factura_relacionada_con_la_operación', 'N/A'),
            '${informe_si_cuenta_con_fotografias_videos_o_informacion_adicion}': data.get('informe_si_cuenta_con_fotografias_videos_o_informacion_adicion', 'N/A'),
            '${comentarios}': data.get('comentarios', 'N/A')
        }

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for template in TEMPLATE_FILES:
                try:
                    doc_buffer = generate_single_document(template, template_root, replacements, user_image_path, data)
                    base = os.path.splitext(template)[0]
                    rfc = data.get('r_f_c', 'N/A')
                    output_filename = f"{sanitize_filename(base)}_{sanitize_filename(descripcion_servicio)}_{sanitize_filename(base)}_{numero_de_contrato_unico}_{sanitize_filename(rfc)}.docx"
                    try:
                        output_path = GENERATED_DOCS / output_filename
                        output_path.parent.mkdir(parents=True, exist_ok=True)
                        with open(output_path, "wb") as f:
                            f.write(doc_buffer.getvalue())
                    except Exception as e:
                        logger.debug("No se pudo guardar en disco (ok en serverless): %s", e)
                    zip_file.writestr(output_filename, doc_buffer.getvalue())
                except Exception as e:
                    logger.exception("Error generando documento %s: %s", template, e)
                    continue

        zip_buffer.seek(0)
        final_zip_name = f"{sanitize_filename(descripcion_servicio)}_{numero_de_contrato_unico}_{sanitize_filename(data.get('r_f_c', 'N/A'))}.zip"

        try:
            tmp_zip = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
            tmp_zip.write(zip_buffer.getvalue())
            tmp_zip.flush()
            tmp_zip_name = tmp_zip.name
            tmp_zip.close()
            try:
                pathlib.Path(GENERATED_ZIPS).mkdir(parents=True, exist_ok=True)
                dest = pathlib.Path(GENERATED_ZIPS) / final_zip_name
                with open(tmp_zip_name, "rb") as s, open(dest, "wb") as d:
                    d.write(s.read())
            except Exception:
                pass
        except Exception as e:
            logger.warning("No se pudo crear zip temporal en disco: %s", e)
            tmp_zip_name = None

        zip_buffer.seek(0)
        upload_result = authenticate_and_upload_to_drive(final_zip_name, zip_buffer)
        logger.info("GOOGLE DRIVE: %s", upload_result.get("message"))

        zip_buffer.seek(0)
        response = send_file(zip_buffer, mimetype='application/zip', as_attachment=True, download_name=final_zip_name)

        try:
            if tmp_zip_name and os.path.exists(tmp_zip_name):
                os.remove(tmp_zip_name)
        except Exception:
            pass

        return response

    except Exception as e:
        logger.exception("Error interno en endpoint: %s", e)
        return jsonify({"error": f"Error interno: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(debug=True)
