import os
import io
import zipfile
import random
import json
import base64
import logging
import tempfile
import pathlib
import re
import time
import boto3

from botocore.exceptions import ClientError, ParamValidationError
from mangum import Mangum
from googleapiclient.errors import HttpError

from flask import (
    Flask, request, send_file, jsonify, abort, Response, make_response
)
from flask_cors import CORS

from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge

from docx import Document
from docx.shared import Inches

from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

from datetime import datetime, timedelta, timezone
from collections import defaultdict
from functools import wraps

# --------------------------
# SES Imports
# --------------------------
import email.encoders
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication

# --------------------------
# App Configuración Base
# --------------------------
app = Flask(__name__)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("generate-word-app")

# Tamaños máximos configurables
MAX_CONTENT_LENGTH_BYTES = int(os.getenv("MAX_CONTENT_LENGTH_BYTES", 20 * 1024 * 1024))
ALLOWED_IMAGE_EXT = {".png", ".jpg", ".jpeg"}

app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH_BYTES
app.config["JSON_SORT_KEYS"] = False

# Carpeta de plantillas
TEMPLATE_FOLDER_NAME = os.getenv("TEMPLATE_FOLDER_NAME", "template_word")
TEMPLATE_FOLDER = pathlib.Path(__file__).parent / TEMPLATE_FOLDER_NAME

# CORS (Relevante para el punto 1: Frontend/Amplify)
allowed_origins = os.getenv("ALLOWED_ORIGINS", "")
if allowed_origins:
    origins = [o.strip() for o in allowed_origins.split(",") if o.strip()]
    CORS(app, origins=origins)
else:
    CORS(app, resources={r"/generate-word": {"origins": []}})

# --------------------------
# AWS COGNITO Config (Control de acceso para contactos.html)
# --------------------------
COGNITO_USER_POOL_ID = os.getenv("us-east-2_sWJSQ4mrD")
COGNITO_CLIENT_ID = os.getenv("4hdk0upvrq9h0p9s8v5ib1th48")
COGNITO_REGION = os.getenv("us-east-2")

if not COGNITO_USER_POOL_ID or not COGNITO_CLIENT_ID:
    logger.error("COGNITO_USER_POOL_ID o COGNITO_CLIENT_ID no configurados.")

try:
    # Operación: Inicialización del cliente para interactuar con Cognito IDP
    cognito_client = boto3.client("cognito-idp", region_name=COGNITO_REGION)
except Exception as e:
    logger.error(f"Error al inicializar cliente Cognito: {e}")

# --------------------------
# AWS SES Config (Servicio de envío de correo)
# --------------------------
SES_SOURCE_EMAIL = os.getenv("SES_SOURCE_EMAIL", "noreply@example.com")
DESTINATION_EMAIL = os.getenv("DESTINATION_EMAIL")

# --------------------------
# Rate Limiting (Protección para la ruta /login)
# --------------------------
MAX_ATTEMPTS = 5
BLOCK_TIME_SECONDS = 300

login_attempts = defaultdict(
    lambda: {"count": 0, "last_attempt": 0, "blocked_until": 0}
)

# --------------------------
# Mapeo de Servicios (Relaciona la entrada del formulario con carpetas de plantillas)
# --------------------------
SERVICIO_TO_DIR = {
    "Servicios de construccion de unidades unifamiliares": "construccion_unifamiliar",
    "Servicios de reparacion o ampliacion o remodelacion de viviendas unifamiliares": "reparacion_remodelacion_unifamiliar",
    "Servicio de remodelacion general de viviendas unifamiliares": "remodelacion_general",
    "Servicios de reparacion de casas moviles en el sitio": "reparacion_casas_moviles",
    "Servicios de construccion y reparacion de patios y terrazas": "patios_terrazas",
    "Servico de reparacion por daños ocasionados por fuego de viviendas unifamiliares": "reparacion_por_fuego",
    "Servicio de construccion de casas unifamiliares nuevas": "construccion_unifamiliar_nueva",
    "Servicio de instalacion de casas unifamiliares prefabricadas": "instalacion_prefabricadas",
    "Servicio de construccion de casas en la ciudad o casas jardin unifamiliares nuevas": "construccion_casas_ciudad_jardin",
    "Dasarrollo urbano": "desarrollo_urbano",
    "Servicio de planificacion de la ordenacion urbana": "planificacion_ordenacion_urbana",
    "Servicio de administracion de tierras urbanas": "administracion_tierras_urbanas",
    "Servicio de programacion de inversiones urbanas": "programacion_inversiones_urbanas",
    "Servicio de reestructuracion de barrios marginales": "reestructuracion_barrios_marginales",
    "Servicios de alumbrado urbano": "alumbrado_urbano",
    "Servicios de control o regulacion del desarrollo urbano": "control_desarrollo_urbano",
    "Servicios de estandares o regulacion de edificios urbanos": "estandares_regulacion_edificios",
    "Servicios comunitarios urbanos": "comunitarios_urbanos",
    "Servicios de administracion o gestion de proyectos o programas urbanos": "gestion_proyectos_programas_urbanos",
    "Ingenieria civil": "ingenieria_civil",
    "Ingenieria de carreteras": "ingenieria_carreteras",
    "Ingenieria de infraestructura de instalaciones o fabricas": "infraestructura_instalaciones_fabricas",
    "Servicios de mantenimiento e instalacion de equipo pesado": "mantenimiento_instalacion_equipo_pesado",
    "Servicio de mantenimiento y reparacion de equipo pesado": "mantenimiento_reparacion_equipo_pesado",
}

TEMPLATE_FILES = [
    "plantilla_solicitud.docx",
    "2.docx",
    "3.docx",
    "4.docx",
    "1.docx",
]

# --------------------------
# Utilidades
# --------------------------

# Operación: Normaliza y acorta el nombre del archivo para seguridad.
def sanitize_filename(name: str) -> str:
    name = secure_filename(name)
    name = re.sub(r"[_]{2,}", "_", name)
    return name[:200]

# Operación: Realiza la sustitución de marcadores (${KEY}) por valores en el documento (Punto 3).
def replace_text_in_document(document, replacements):
    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))

# Operación: Carga la plantilla, sustituye el texto, inserta la imagen de firma y guarda en un buffer.
def generate_single_document(template_filename, template_root, replacements, user_image_path=None, data=None):
    template_path = template_root / template_filename

    if not template_path.exists():
        raise FileNotFoundError(f"Plantilla '{template_filename}' no encontrada.")

    document = Document(template_path)
    replace_text_in_document(document, replacements)

    if user_image_path and os.path.exists(user_image_path):
        try:
            document.add_paragraph()
            document.add_paragraph(
                data.get(
                    "nombre_completo_de_la_persona_que_firma_la_solicitud",
                    "N/A",
                )
            )
            document.add_picture(user_image_path, width=Inches(2.5))
        except Exception as ex:
            logger.warning("No se pudo insertar imagen: %s", ex)
            document.add_paragraph("⚠ No se pudo insertar la imagen del usuario.")
    else:
        document.add_paragraph("⚠ Imagen de firma no encontrada.")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# Operación: Obtiene un secreto (usado para Google Service Account) de AWS Secrets Manager.
def get_secret_value(secret_name):
    if not secret_name:
        raise EnvironmentError("Secret name not provided.")

    try:
        client = boto3.client("secretsmanager")
        resp = client.get_secret_value(SecretId=secret_name)

        if "SecretString" in resp and resp["SecretString"]:
            return resp["SecretString"]

        return resp["SecretBinary"]

    except ClientError as e:
        logger.exception("No se pudo obtener secreto %s: %s", secret_name, e)
        raise

# --------------------------
# Google Drive Upload (Paso 5: Envía carpeta comprimida a Google Drive)
# --------------------------

# Operación: Carga las credenciales de la cuenta de servicio de Google desde variables de entorno o Secrets Manager.
def _load_service_account_info():
    sm_name = os.getenv("GOOGLE_SERVICE_ACCOUNT_SECRET_NAME")

    if sm_name:
        raw = get_secret_value(sm_name)
        return json.loads(raw)

    raw_json = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
    if raw_json:
        return json.loads(raw_json)

    b64 = os.getenv("GOOGLE_SERVICE_ACCOUNT_BASE64")
    if b64:
        decoded = base64.b64decode(b64)
        return json.loads(decoded)

    raise EnvironmentError("Falta GOOGLE_SERVICE_ACCOUNT")

# Operación: Autentica usando Service Account y sube el archivo ZIP a la carpeta configurada de Drive.
def authenticate_and_upload_to_drive(file_name, zip_buffer):
    if os.getenv("DISABLE_DRIVE_UPLOAD", "0") == "1":
        logger.info("Subida a Drive deshabilitada.")
        return {"success": True, "message": "Subida deshabilitada"}

    try:
        service_account_info = _load_service_account_info()
    except Exception as e:
        logger.exception("No se pudo cargar service account: %s", e)
        return {"success": False, "message": str(e)}

    scopes = ["https://www.googleapis.com/auth/drive.file"]

    try:
        creds = Credentials.from_service_account_info(
            service_account_info,
            scopes=scopes
        )
        service = build("drive", "v3", credentials=creds, cache_discovery=False)

        file_metadata = {
            "name": sanitize_filename(file_name),
            "mimeType": "application/zip",
        }

        drive_folder = os.getenv("DRIVE_FOLDER_ID")
        if drive_folder:
            file_metadata["parents"] = [drive_folder]

        zip_buffer.seek(0)
        media = MediaIoBaseUpload(zip_buffer, mimetype="application/zip", resumable=False)

        max_retries = 3
        for attempt in range(1, max_retries + 1):
            try:
                uploaded = (
                    service.files()
                    .create(body=file_metadata, media_body=media, fields="id")
                    .execute()
                )
                return {
                    "success": True,
                    "message": f"Archivo subido. ID: {uploaded.get('id')}",
                }
            except HttpError as he:
                logger.warning("Error en intento %d: %s", attempt, he)
                if attempt == max_retries:
                    return {"success": False, "message": str(he)}
                time.sleep(2**attempt)

    except Exception as e:
        logger.exception("Error general Drive: %s", e)
        return {"success": False, "message": str(e)}

# --------------------------
# Envío Ses (Paso 6: Envía correo al cliente)
# --------------------------

# Operación: Utiliza AWS SES para enviar el archivo ZIP adjunto a la dirección de correo configurada (DESTINATION_EMAIL).
def send_email_with_attachment(zip_buffer, filename, recipient_email):
    if not recipient_email:
        logger.warning("DESTINATION_EMAIL no configurado.")
        return {"success": False, "message": "DESTINATION_EMAIL no configurado"}

    logger.info(f"Enviando correo a {recipient_email}...")

    msg = MIMEMultipart()
    msg["Subject"] = f"Documentos Generados: {filename.replace('.zip', '')}"
    msg["From"] = SES_SOURCE_EMAIL
    msg["To"] = recipient_email

    msg.attach(MIMEApplication("Adjunto archivo ZIP.", _subtype="plain"))

    zip_buffer.seek(0)
    att = MIMEApplication(zip_buffer.read(), _subtype="zip")
    att.add_header("Content-Disposition", "attachment", filename=filename)
    email.encoders.encode_base64(att)
    msg.attach(att)
    zip_buffer.seek(0)

    try:
        ses_client = boto3.client("ses")
        response = ses_client.send_raw_email(
            Source=SES_SOURCE_EMAIL,
            Destinations=[recipient_email],
            RawMessage={"Data": msg.as_string()},
        )
        return {"success": True, "message": f"Correo enviado. ID: {response['MessageId']}"}

    except ClientError as e:
        error_message = e.response["Error"]["Message"]
        logger.error("Error SES: %s", error_message)
        return {"success": False, "message": f"Error SES: {error_message}"}

    except Exception as e:
        logger.exception("Error general SES")
        return {"success": False, "message": f"Error general: {e}"}

# --------------------------
# Decorador Token Cognito (Controla el acceso al formulario en contactos.html)
# --------------------------

# Operación: Revisa si la IP ha excedido el límite de intentos de login y si está bloqueada.
def check_rate_limit(ip):
    entry = login_attempts[ip]
    now = time.time()

    if entry["blocked_until"] > now:
        return False, int(entry["blocked_until"] - now)

    return True, 0

# Operación: Registra un intento fallido y bloquea al usuario si excede el límite.
def record_failed_attempt(ip):
    entry = login_attempts[ip]
    now = time.time()

    entry["count"] += 1
    entry["last_attempt"] = now

    if entry["count"] >= MAX_ATTEMPTS:
        entry["blocked_until"] = now + BLOCK_TIME_SECONDS

# Operación: Resetea el contador de intentos fallidos.
def reset_attempts(ip):
    login_attempts[ip] = {
        "count": 0,
        "last_attempt": 0,
        "blocked_until": 0,
    }

# Decorador: Valida el token JWT de Cognito en la cabecera Authorization para asegurar accesos controlados.
def token_required(required_role="user"):
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            token = None

            auth_header = request.headers.get("Authorization")
            if auth_header and auth_header.startswith("Bearer "):
                token = auth_header.split(" ")[1]

            if not token:
                return jsonify({"error": "Token faltante"}), 401

            try:
                # Operación: Llama a Cognito get_user para validar el token y obtener la identidad del usuario.
                response = cognito_client.get_user(AccessToken=token)
                current_user = response["Username"]

                user_role = "user"  # por defecto

                if user_role != required_role:
                    return jsonify({"error": "Sin permisos"}), 403

            except ClientError as e:
                error_code = e.response["Error"]["Code"]

                if error_code in [
                    "NotAuthorizedException",
                    "ExpiredTokenException",
                    "InvalidParameterException",
                ]:
                    return jsonify({"error": "Token inválido"}), 401

                return jsonify({"error": "Error interno de autenticación"}), 500

            except Exception as e:
                logger.error("Error en token_required: %s", e)
                return jsonify({"error": "Error interno del servidor"}), 500

            return f(current_user, *args, **kwargs)

        return decorated

    return decorator

# --------------------------
# Login Cognito (Ruta expuesta en contactos.html)
# --------------------------
@app.route("/login", methods=["POST"])
def login():
    ip = request.remote_addr
    allowed, wait = check_rate_limit(ip)

    if not allowed:
        return jsonify({"error": f"Demasiados intentos. Espera {wait}s"}), 429

    data = request.get_json()
    username = data.get("username")
    password = data.get("password")

    if not username or not password:
        return jsonify({"error": "Faltan datos"}), 400

    if not COGNITO_USER_POOL_ID or not COGNITO_CLIENT_ID:
        return jsonify({"error": "Falta configuración de Cognito"}), 500

    try:
        # Operación: Inicia el flujo de autenticación de Cognito para validar credenciales.
        response = cognito_client.initiate_auth(
            AuthFlow="USER_PASSWORD_AUTH",
            AuthParameters={
                "USERNAME": username,
                "PASSWORD": password,
            },
            ClientId=COGNITO_CLIENT_ID,
        )

        if "AuthenticationResult" in response:
            reset_attempts(ip)

            return jsonify({
                "message": "Login exitoso",
                "token": response["AuthenticationResult"]["AccessToken"],
                "id_token": response["AuthenticationResult"]["IdToken"],
            }), 200

        elif response.get("ChallengeName"):
            return jsonify({
                "error": "Se requiere un desafío adicional"
            }), 403

    except ClientError as e:
        error = e.response["Error"]["Code"]

        if error in ["NotAuthorizedException", "UserNotFoundException"]:
            record_failed_attempt(ip)
            return jsonify({"error": "Credenciales inválidas"}), 401

        if error == "TooManyRequestsException":
            return jsonify({"error": "Demasiadas solicitudes"}), 429

        logger.error("Error Cognito login: %s", error)
        return jsonify({"error": "Error interno Cognito"}), 500

    except Exception as e:
        logger.exception("Error general login")
        return jsonify({"error": "Error interno del servidor"}), 500


# --------------------------
# generate-word (Ruta principal protegida por token, implementa flujo de documentos)
# --------------------------
@app.route("/generate-word", methods=["POST"])
@token_required(required_role="user") # Acceso controlado
def generate_word(current_user):
    logger.info(f"Solicitud de generación: {current_user}")

    user_image_path = None

    try:
        data = request.form.to_dict()
        uploaded_image = request.files.get("imagen_usuario")

        if not data:
            return jsonify({"error": "No data"}), 400

        servicio = data.get("servicio")
        carpeta = SERVICIO_TO_DIR.get(servicio)

        if not carpeta:
            return jsonify({"error": f"Servicio desconocido: {servicio}"}), 404

        template_root = TEMPLATE_FOLDER / carpeta
        if not template_root.is_dir():
            return jsonify({"error": f"Carpeta no existe: {template_root}"}), 404

        # Imagen temporal
        if uploaded_image and uploaded_image.filename:
            # ... Lógica para sanear, validar tamaño y guardar imagen temporalmente ...
            filename = sanitize_filename(uploaded_image.filename)
            ext = pathlib.Path(filename).suffix.lower()

            if ext not in ALLOWED_IMAGE_EXT:
                return jsonify({"error": "Tipo de imagen no permitido"}), 400

            tmp_dir = pathlib.Path(tempfile.gettempdir()) / f"upload_{os.getpid()}"
            tmp_dir.mkdir(exist_ok=True)

            user_image_path = str(tmp_dir / filename)

            MAX_IMAGE_BYTES = int(os.getenv("MAX_IMAGE_BYTES", 5 * 1024 * 1024))
            uploaded_image.stream.seek(0, io.SEEK_END)

            size = uploaded_image.stream.tell()
            uploaded_image.stream.seek(0)

            if size > MAX_IMAGE_BYTES:
                return jsonify({"error": "Imagen demasiado grande"}), 413

            uploaded_image.save(user_image_path)

        # Número aleatorio (para el nombre del documento)
        numero = "".join([str(random.randint(0, 9)) for _ in range(18)])

        # Mapeo de Reemplazos
        replacements = {
            '${descripcion_del_servicio}': servicio,
            '${razon_social}': data.get('razon_social', 'N/A'),
            '${r_f_c}': data.get('r_f_c', 'N/A'),
            '${domicilio_del_cliente}': data.get('domicilio_del_cliente', 'N/A'),
            '${telefono_del_cliente}': data.get('telefono_del_cliente', 'N/A'),
            '${correo_del_cliente}': data.get('correo_del_cliente', 'N/A'),
            '${fecha_de_inicio_del_servicio}': data.get('fecha_de_inicio_del_servicio', 'N/A'),
            '${fecha_de_conclusion_del_servicio}': data.get('fecha_de_conclusion_del_servicio', 'N/A'),
            '${monto_de_la_operacion_Sin_IVA}': data.get('monto_de_la_operacion_in_iva', 'N/A'),
            '${forma_de_pago}': data.get('forma_de_pago', 'N/A'),
            '${cantidad}': data.get('cantidad', 'N/A'),
            '${unidad}': data.get('unidad', 'N/A'),
            '${numero_de_contrato}': numero,
            '${fecha_de_operación}': data.get('fecha_de_operacion', 'N/A'),
            '${nombre_completo_de_la_persona_que_firma_la_solicitud}': data.get('nombre_completo_de_la_persona_que_firma_la_solicitud', 'N/A'),
            '${cargo_de_la_persona_que_firma_la_solicitud}': data.get('cargo_de_la_persona_que_firma_la_solicitud', 'N/A'),
            '${factura_relacionada_con_la_operación}': data.get('factura_relacionada_con_la_operación', 'N/A'),
            '${informe_si_cuenta_con_fotografias_videos_o_informacion_adicion}': data.get('informe_si_cuenta_con_fotografias_videos_o_informacion_adicion', 'N/A'),
            '${comentarios}': data.get('comentarios', 'N/A')
        }

        # Inicializar ZIP buffer (Paso 4: Comprime)
        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for template_filename in TEMPLATE_FILES:
                # Operación: Genera documento individual, reemplaza texto e inserta imagen (Paso 3)
                buffer = generate_single_document(
                    template_filename,
                    template_root,
                    replacements,
                    user_image_path=user_image_path,
                    data=data
                )

                # Operación: Añade el documento recién generado al archivo ZIP
                zipf.writestr(
                    f"{template_filename}",
                    buffer.getvalue()
                )

        zip_buffer.seek(0)
        final_filename = f"documentos_{numero}.zip"

        # Envío por correo (Paso 6)
        send_email_with_attachment(zip_buffer, final_filename, DESTINATION_EMAIL)

        # Subir a drive (Paso 5)
        authenticate_and_upload_to_drive(final_filename, zip_buffer)

        zip_buffer.seek(0)

        # Respuesta: Retorna el archivo ZIP al cliente del frontend (descarga final)
        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name=final_filename,
            mimetype="application/zip",
        )

    except RequestEntityTooLarge:
        return jsonify({"error": "Archivo demasiado grande"}), 413

    except Exception as e:
        logger.exception("Error general en generate-word")
        return jsonify({"error": "Error interno del servidor"}), 500


# --------------------------
# Handler Lambda
# --------------------------
# Operación: Adapta la aplicación Flask para ser utilizada como un handler de AWS Lambda/API Gateway
handler = Mangum(app)